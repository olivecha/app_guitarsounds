import os
import datetime
from inspect import getdoc
import numpy as np
from docx import Document
from docx.shared import Inches
import streamlit as st
import matplotlib.pyplot as plt
from PIL import Image
from guitarsounds.analysis import Plot, Signal, Sound, SoundPack
from defined_analyses import all_report_headers

def get_file_number(filename):
    """ get the number in a filename """
    numbers = []
    for char in filename:
        try:
            numbers.append(int(char))
        except:
            pass 
    numbers = [str(num) for num in numbers]
    return int(''.join(numbers))

def get_temp_sound_number():
    """ 
    Get the current temp sound number 
    """
    numbers = []
    for filename in os.listdir('temp_sounds'):
        numbers.append(get_file_number(filename))

    if len(numbers) == 0:
        return 0
    else:
        return np.sort(numbers)[-1] + 1

def get_cached_next_number(session_state):
    """ 
    Find the next sound key to cache 
    """
    numbers = list(session_state['sounds_cache'].keys())
    numbers = [int(num) for num in numbers]
    if len(numbers) > 0:
        return np.sort(numbers)[-1] +1 
    else:
        return 0

def remove_cached_sound(session_state, sound_key):
    """ 
    remove a sound from the cached sounds 
    """
    session_state['sounds_cache'].pop(sound_key)

def remove_log_ticks(fig):
    """
    Make the ticks in log scale plots linear so its less confusing
    """
    for ax in fig.axes:
        labels = ax.get_xticklabels()
        xmin, xmax = ax.get_xlim()
        print(xmin)
        new_labels = []
        for l in labels:
            try:
                _, number, exp = l.get_text().split('{')
                number = number.replace('^', "")
                exp = exp.split('}')[0]
                value = int(number) ** int(exp)
                new_labels.append(value)
            except ValueError:
                nstr = l.get_text()
                nstr = nstr.replace('−', '-')
                if '.' in nstr:
                    new_labels.append(float(nstr))
                else:
                    new_labels.append(int(nstr))
        new_labels = [l for l in new_labels if l > xmin]
        new_labels = [l for l in new_labels if l < xmax]
        ax.set_xticks(new_labels)
        ax.set_xticklabels(new_labels)

def create_figure(analysis, key, *args):
    """ Run a plotting fonction but include calls to streamlit """
    fig, ax   = plt.subplots(figsize=(7, 4.5))
    plt.sca(ax)
    analysis(*args)
    fig = plt.gcf()
    #TODO: make less sketch
    try:
        remove_log_ticks(fig)
    except:
        pass
    fig.savefig(os.path.join('figure_cache',key), dpi=200)

def generate_report(report_analyses):
    """Génère un rapport d'analyse en format word"""
    document = Document()
    today = datetime.date.today()
    document.add_heading(f"Rapport d'analyse comparative de sons de guitare", level=0)
    document.add_heading(f'({today.day}/{today.month}/{today.year})', level=2)
    for analysis in report_analyses:
        if report_analyses[analysis]:
            document.add_paragraph(all_report_headers[analysis], style='Heading 2')
            document.add_picture(f'figure_cache/{analysis}.png',width=Inches(5))
        document.save('report.docx')


