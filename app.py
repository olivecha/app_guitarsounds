import os, io
import pickle
import datetime
import numpy as np
import streamlit as st
import defined_analyses
import matplotlib.pyplot as plt
from PIL import Image
from docx import Document
from docx.shared import Inches
from pydub import AudioSegment
from tempfile import NamedTemporaryFile
from audiorecorder import audiorecorder
from guitarsounds.analysis import SoundPack
from app_utils import get_cached_next_number, display_norm_help
from app_utils import remove_cached_sound, audioseg2guitarsound
from app_utils import mpl2pil, remove_log_ticks, remove_log_ticksY


def save_state(obj):
    """
    Takes the st.session state as an argument 
    and saves the user defined attributes to a pickle file
    (This is used to reload the application from a previous state
     manually to try and test the analyses without running the
     graphical interface, work in progress)
    """
    session_state_keys = ['upload_status',
                          'analysis_menu',
                          'report_analyses',
                          #'cached_images',
                          "report_file",
                          'loop_no',
                          'cached_logo',
                          'number_of_sounds',
                          'current_sounds',
                          'cached_sounds',
                          'reference_recording',
                          'norm_bool',
                          'norm_toggle']
    print('Saving state')
    state = {}
    for ky in session_state_keys:
        state[ky] = obj[ky]
    with open('/Users/olivier/Desktop/session_state.pkl', 'wb') as pfile:
        pickle.dump(state, pfile)

# """
# Functions modifying the session state 
# Need to be defined in this file
# """
def update_file_load_status():
    """ function for upload button call_back """
    st.session_state['upload_status'] = True

def reset_analyses(analysis_names):
    """ 
    Function that resets the value of the 
    analyses to their default state 
    """
    st.session_state['analysis_menu'] = {name:None for name in analysis_names}
    st.session_state['report_analyses'] = {name:False for name in analysis_names}
    st.session_state['cached_images'] = {name:None for name in analysis_names}

def set_state(menu, analysis, state):
    """ set the session state of a menu to organise outputs """
    st.session_state[menu][analysis] = state 

def generate_figure_and_set_state(analysis_call, key, sound):
    """
    Function to alter the session state depending on the analysis call
    """
    # Case for the listen frequency bins function
    if key == 'listenband':
        st.session_state['analysis_menu'][key] = 'listen'
    # Case for interactive analyses
    elif key in ['signal', 'envelope', 'logenv', 'fft', 'mfft']:
        st.session_state['analysis_menu'][key] = 'call'
    # Default
    else:
        image = create_figure2(analysis_call, key, sound) 
        st.session_state['cached_images'][key] = image
        st.session_state['analysis_menu'][key] = 'figure'

def create_figure3(analysis, key, *args):
    """ Create figure and updates state """
    fig, ax   = plt.subplots(figsize=(7, 4.5))
    plt.sca(ax)
    analysis(*args)
    fig = plt.gcf()
    remove_log_ticks(fig)
    remove_log_ticksY(fig)
    image = mpl2pil(fig)
    st.session_state['cached_images'][key] = image

def create_figure2(analysis, key, *args):
    """ Run a plotting fonction but include calls to streamlit """
    fig, ax   = plt.subplots(figsize=(7, 4.5))
    plt.sca(ax)
    analysis(*args)
    fig = plt.gcf()
    remove_log_ticks(fig)
    remove_log_ticksY(fig)
    image = mpl2pil(fig)
    return image

def generate_report(report_analyses, cached_images):
    """Génère un rapport d'analyse en format word"""
    document = Document()
    today = datetime.date.today()
    document.add_heading(f"Rapport d'analyse comparative de sons de guitare", level=0)
    document.add_heading(f'({today.day}/{today.month}/{today.year})', level=2)
    for analysis in report_analyses:
        if report_analyses[analysis]:
            document.add_paragraph(defined_analyses.all_report_headers[analysis], style='Heading 2')
            tempFile = io.BytesIO()
            image = cached_images[analysis]
            image.save(tempFile, format="PNG")
            document.add_picture(tempFile, width=Inches(5))
        tempDoc = io.BytesIO()
        document.save(tempDoc)
        tempDoc.seek(0)
        st.session_state["report_file"] = tempDoc

def variable_signal_context(sound):
    """
    Menu for the variable time signal plot
    """
    c = st.container()
    colmin, colmax, colgo = c.columns([2, 2, 1])
    lower_bound = colmin.number_input(label="Temps min", 
                                      min_value=0.0)
    upper_bound = colmax.number_input(label="Temps max", 
                                      max_value=sound.signal.time()[-1],
                                      value=sound.signal.time()[-1])
    colgo.write(" ")
    colgo.write(" ")
    colgo.button(label="Actualiser", 
                 key='Actualiser signal',
                 on_click=create_figure3, 
                 args=(defined_analyses.variable_signal_plot, 
                       'signal', 
                        sound,
                        lower_bound,
                        upper_bound,
                        c))
    # Change the session state if there is no image
    if st.session_state['cached_images']['signal'] is None:
        fun = defined_analyses.plot_with_sound(defined_analyses.Plot.signal)
        img = create_figure2(fun,
                             'signal',
                             sound)
        st.session_state['cached_images']['signal'] = img

def variable_envelope_context(sound):
    """
    Menu for the variable time envelope plot
    """
    c = st.container()
    colmin, colmax, colgo = c.columns([2, 2, 1])
    lower_bound = colmin.number_input(label="Temps min", 
                                      key='envelope min',
                                      min_value=0.0)
    upper_bound = colmax.number_input(label="Temps max", 
                                      key='envelope max',
                                      max_value=sound.signal.time()[-1],
                                      value=sound.signal.time()[-1])
    colgo.write(" ")
    colgo.write(" ")
    colgo.button(label="Actualiser", 
                 key='Actialiser envelope',
                 on_click=create_figure3, 
                 args=(defined_analyses.variable_envelope_plot, 
                       'envelope', 
                        sound,
                        lower_bound,
                        upper_bound,
                        c))
    if st.session_state['cached_images']['envelope'] is None:
        print("Generating first envelope figure")
        fun = defined_analyses.plot_with_sound(defined_analyses.Plot.envelope)
        img = create_figure2(fun,
                             'envelope',
                             sound)
        st.session_state['cached_images']['envelope'] = img

def variable_logenv_context(sound):
    """
    Menu for the variable time log envelope plot
    """
    c = st.container()
    colmin, colmax, colgo = c.columns([2, 2, 1])
    lower_bound = colmin.number_input(label="Temps min", 
                                      key='logenv min',
                                      min_value=0.0)
    upper_bound = colmax.number_input(label="Temps max", 
                                      key='logenv max',
                                      max_value=sound.signal.time()[-1],
                                      value=sound.signal.time()[-1])
    colgo.write(" ")
    colgo.write(" ")
    colgo.button(label="Actualiser", 
                 key='Actualiser logenv',
                 on_click=create_figure3, 
                 args=(defined_analyses.variable_logenv_plot, 
                       'logenv', 
                        sound,
                        lower_bound,
                        upper_bound,
                        c))

    if st.session_state['cached_images']['logenv'] is None:
        fun = defined_analyses.plot_with_sound(defined_analyses.Plot.log_envelope)
        img = create_figure2(fun,
                             'logenv',
                             sound)

def variable_fft_context(sound):
    """
    Menu for the variable time log envelope plot
    """
    c = st.container()
    colmin, colmax, colgo = c.columns([2, 2, 1])
    lower_bound = colmin.number_input(label="Fréquence min", 
                                      min_value=0)
    upper_bound = colmax.number_input(label="Fréquence max", 
                                      max_value=3000,
                                      value=2000)
    colgo.write(" ")
    colgo.write(" ")
    colgo.button(label="Actualiser", 
                 on_click=create_figure3, 
                 key='Actualiser fft',
                 args=(defined_analyses.variable_fft_plot, 
                       'fft', 
                        sound,
                        lower_bound,
                        upper_bound,
                        c))

    if st.session_state['cached_images']['fft'] is None:
        fun = defined_analyses.plot_with_sound(defined_analyses.Plot.fft)
        img = create_figure2(fun,
                             'fft',
                             sound)
        st.session_state['cached_images']['fft'] = img

def variable_dual_fft_context(sounds):
    """
    Menu for the dual fft plot
    """
    c = st.container()
    colmin, colmax, colgo = c.columns([2, 2, 1])
    lower_bound = colmin.number_input(label="Fréquence min", 
                                      min_value=0)
    upper_bound = colmax.number_input(label="Fréquence max", 
                                      max_value=3000,
                                      value=2000)
    colgo.write(" ")
    colgo.write(" ")
    colgo.button(label="Actualiser", 
                 on_click=create_figure3, 
                 key='Actualiser fft',
                 args=(defined_analyses.variable_dual_fft_plot, 
                       'mfft', 
                        sounds,
                        lower_bound,
                        upper_bound,
                        c))

    if st.session_state['cached_images']['mfft'] is None:
        print('Image not yet cached in session state')
        fun = defined_analyses.plot_with_soundpack(defined_analyses.Plot.fft)
        img = create_figure2(fun,
                             'mfft',
                             sounds)
        st.session_state['cached_images']['mfft'] = img

# """
# Update the defined analyses with the locally defined functions
# """
defined_analyses.single_sound_analysis_functions['signal'] = variable_signal_context
defined_analyses.single_sound_analysis_functions['envelope'] = variable_envelope_context
defined_analyses.single_sound_analysis_functions['logenv'] = variable_logenv_context
defined_analyses.single_sound_analysis_functions['fft'] = variable_fft_context
defined_analyses.dual_sound_analysis_functions['mfft'] = variable_dual_fft_context

# """ 
# Session setup
# """

# Print an incremented number at each loop to help
# debugging other print statements
if 'loop_no' not in st.session_state:
    st.session_state['loop_no'] = 0
print('GUI Loop number:', st.session_state['loop_no'])
st.session_state['loop_no'] += 1

# Attribute to store the generated report temporary file
if "report_file" not in st.session_state:
    st.session_state["report_file"] = None

# Flag to not upload the sound for every run
if 'upload_status' not in st.session_state:
    st.session_state['upload_status'] = None

# Store the Bruant logo in an attribute to not read it
# at every loop
if 'cached_logo' not in st.session_state:
    image = Image.open('documentation/figures/logo.png')
    st.session_state['cached_logo'] = image

# Define the number of sounds item
if 'number_of_sounds' not in st.session_state:
    st.session_state['number_of_sounds'] = 0

# Store the values of the sound signals in the session state
if 'current_sounds' not in st.session_state:
    st.session_state['current_sounds'] = [np.zeros(10)]

# Session state data for the analyses
# Define the analysis menu in the session state
if 'analysis_menu' not in st.session_state:
    # Define the analyses attributes
    reset_analyses(defined_analyses.single_sound_analysis_names)
    # Number of souds to keep track of changes
    st.session_state['number_of_sounds'] = 0

# Cache all updated sounds to keep between
# session loops
if 'cached_sounds' not in st.session_state:
    st.session_state['cached_sounds'] = {}

# Create a reference recording to store temporarely
# recorded sounds so we add a new sound only if the
# recorded sound has changed
if 'reference_recording' not in st.session_state:
    st.session_state['reference_recording'] = ''

# Variables to watch if the normalize checkbox was
# toggled which restart the analyses
if 'norm_bool' not in st.session_state:
    # Whether or not multiple sounds should be normalized
    st.session_state['norm_bool'] = False
if 'norm_toggle' not in st.session_state:
    # If the button state changed this becomes true
    st.session_state['norm_toggle'] = False


# Title and logo
col_logo, col_title = st.columns([1, 2])
with col_title:
    st.title("Analyse comparative de sons de guitare (V 2.0)")
with col_logo:
    st.write('')
    st.write('')
    st.image(st.session_state['cached_logo'], use_column_width='always')

# Different tabs of the application
sounds_io, analysis, help_tab, about = st.tabs(["Ajouter des sons", 
                                                "Analyser/Comparer des sons",
                                                "Aide",
                                                "À Propos"])

#"""
# First tab (Input / Output)
#"""
with sounds_io: 
    # Command to save the current session state to recreate it in
    # a testing environment
    #save_state(st.session_state)
    # Recording and uploading columns
    record_col, upload_col = st.columns([1, 1])
    # Add some whitespace
    record_col.write('')
    upload_col.write('')
    # Audio recording sound loading in a context box so its neat
    record_context = record_col.expander("Enregistrer un son")
    with record_context:
        # Custom audio recorder widget
        audio_seg = audiorecorder("Cliquez pour débuter l'enregistrement", "Stop")
        print("Recorded audio segment length:", len(audio_seg))

    # If the recorded sound has changed and is not empty
    if (st.session_state['reference_recording'] != audio_seg.raw_data and 
        len(audio_seg) > 0):
        new_sound = audioseg2guitarsound(audio_seg)
        sound_number = get_cached_next_number(st.session_state)
        st.session_state['cached_sounds'][sound_number] = new_sound
        st.session_state['reference_recording'] = audio_seg.raw_data
        st.warning(("Les sons enregistrés ne sont pas sauvegardés d'une session à l'autre \n"
                    "vous pouvez les télécharger si vous voulez les conserver"), icon="⚠️")
    
    # File uploading (second column)
    expander = upload_col.expander("Téléverser un son")
    uploaded_file = expander.file_uploader("Choose a file", 
                                           on_change=update_file_load_status,
                                           label_visibility='collapsed')
    
    # If a file was uploaded
    if uploaded_file is not None and st.session_state['upload_status']:
        # Always process the audio files with AudioSegment
        ext = uploaded_file.name.split('.')[-1]
        audio_seg = AudioSegment.from_file(uploaded_file, ext)
        new_sound = audioseg2guitarsound(audio_seg)
        sound_number = get_cached_next_number(st.session_state)
        st.session_state['cached_sounds'][sound_number] = new_sound
        # Set the upload status to False so the sound is not added
        # after each loop
        st.session_state['upload_status'] = False

    # Context to choose if the sounds are normalized
    norm_soundpacks = st.checkbox("Normaliser les sons multiples", 
                                      value=False)
    # Update the toggle value if the check was changed
    if norm_soundpacks != st.session_state['norm_bool']:
        st.session_state['norm_toggle'] = True
        st.session_state['norm_bool'] = norm_soundpacks
    # Help message to display the normalization help
    norm_help_expander = st.expander("C'est quoi ?")
    display_norm_help(norm_help_expander)
    
    # Display the list of loaded sounds
    col_name, col_audio, col_dl, col_del= st.columns([2, 2.5, 0.5, 0.5])
    names = []
    for sound_number in st.session_state['cached_sounds']:
        # Get the sound in the session state cache
        sound = st.session_state['cached_sounds'][sound_number]
        # Column to remove loaded sounds
        col_del.write(' ')
        col_del.button(label=':wastebasket:', 
                       key=f'del_{sound_number}',
                       on_click=remove_cached_sound,
                       args=(st.session_state, sound_number))
        # Column to change sound names to annotate the plots
        col_name.write(' ')
        name = col_name.text_input(f'name_{sound_number}', 
                               'Nom du son', 
                               label_visibility='collapsed')
        # Store the names
        names.append(name)
        # Column to listen to the loaded sounds
        col_audio.write(' ')
        col_audio.audio(sound.signal.signal, sample_rate=sound.signal.sr)

        # Column to download the loaded sounds (esp for recorded ones)
        col_dl.write(' ')
        # If a sound name was supplied
        if name != 'Nom du son':
            download_name = f"{name}.wav"
        else:
            download_name = f"Son_{sound_number}.wav"
        # Download widget (in .wav format)
        col_dl.download_button(label=':arrow_down:',
                             data=sound.file_bytes,
                             key=f"download_{sound_number}",
                             file_name=f'{name}.wav',)

# """
# Analysis Tab
# """
with analysis:
    # Get all the sounds in a list
    sounds_list = list(st.session_state['cached_sounds'].values())
    # if there are more than one sound, define a soundpack
    if len(sounds_list) > 1:
        analysis_spack = SoundPack(sounds_list,
                                   names=names,
                                   normalize=norm_soundpacks)
    else:
        analysis_spack = None

    if len(sounds_list) > 0:
        # columns to format the output
        colbut, colhelp, colcheck = st.columns([3.5, 0.5, 1.0])
        # Write the columns headers
        colbut.write('Analyse')
        colhelp.write('Aide')
        colcheck.write('Rapport')

        # Load the analyses when looking only at a single sound
        if len(sounds_list) == 1:
            analysis_names = defined_analyses.single_sound_analysis_names
            analysis_functions = defined_analyses.single_sound_analysis_functions
            analysis_helps = defined_analyses.single_sound_analysis_help
            analysis_figures = defined_analyses.single_sound_analysis_help_figures

        # Load the analyses to compare two sounds
        elif len(sounds_list) == 2:
            analysis_names = defined_analyses.dual_sound_analysis_names
            analysis_functions = defined_analyses.dual_sound_analysis_functions
            analysis_helps = defined_analyses.dual_sound_analysis_help
            analysis_figures = defined_analyses.dual_sound_analysis_help_figures

        # Load the analyses to compare an arbitrary amount of sounds
        elif len(sounds_list) > 2:
            analysis_names = defined_analyses.multi_sound_analysis_names
            analysis_functions = defined_analyses.multi_sound_analysis_functions
            analysis_helps = defined_analyses.multi_sound_analysis_help
            analysis_figures = defined_analyses.multi_sound_analysis_help_figures

        # Restart the analyses if the normalization was toggled
        if st.session_state['norm_toggle']:
            reset_analyses(analysis_names)
            st.session_state['norm_toggle'] = False

        # Update if the number of sounds has changed
        elif len(sounds_list) != st.session_state['number_of_sounds']:
            # If the number of sounds or their value changes 
            # We reset the anaylses to allow displaying outputs
            # for the new sounds
            print((f'The number of sounds changed from {st.session_state["number_of_sounds"]}'
                   f' to {len(sounds_list)}'))
            reset_analyses(analysis_names)
            # Update the stored number of sounds
            st.session_state['number_of_sounds'] = len(st.session_state['cached_sounds'])

        # For each analysis defined according to the number of uploaded sounds
        for analysis in analysis_names:
            # Button to run the analysis,
            # Button to display help
            # Check box to add the result to the report
            colbut, colhelp, colcheck = st.columns([3.5, 0.5, 1.0])

            # Buttons with behaviours derived from the analysis functions
            if analysis_spack is None:
                # If there is only one sound the analysis is called on the
                # first element of the sound list
                colbut.button(label=analysis_names[analysis],
                              on_click=generate_figure_and_set_state,
                              use_container_width=True,
                              args=(analysis_functions[analysis], 
                                    analysis, 
                                    sounds_list[0]))
            # Case for soundpacks (>= 2 sounds)
            else:
                # Button to run the analysis on the soundpack
                colbut.button(label=analysis_names[analysis],
                              on_click=generate_figure_and_set_state,
                              use_container_width=True,
                              args=(analysis_functions[analysis], 
                                    analysis, 
                                    analysis_spack))

            # Help button for the current analysis
            colhelp.button(label=':question:',
                           key=analysis,
                           on_click=set_state,
                           args=('analysis_menu', analysis, 'help'))

            # No checkbox for audio output analysis
            # (Cannot put audio in the report)
            if analysis == 'listenband':
                colcheck.write('')
                st.session_state['report_analyses'][analysis] = False

            # For every other analysis
            else:
                # If the analysis was called add a checked checkbox in the report option
                # column
                if st.session_state['analysis_menu'][analysis] in ('figure', 'call'):
                    check = colcheck.checkbox(label=analysis,
                                              value=True, 
                                              label_visibility='collapsed')
                    # Store the check value in the session state
                    st.session_state['report_analyses'][analysis] = check
                # Put white space if the analysis was not called
                else:
                    colcheck.write('')
                    st.session_state['report_analyses'][analysis] = False

            # Get the image produced by matplotlib and display it
            if st.session_state['analysis_menu'][analysis] == 'figure':
                image = st.session_state['cached_images'][analysis]
                st.empty().image(image)

            # Special case for interactive analyses where the function
            # is called on each loop to update the figure
            elif st.session_state['analysis_menu'][analysis] == 'call':
                if analysis in ['mfft']:
                    image = analysis_functions[analysis](analysis_spack)
                else:
                    image = analysis_functions[analysis](sounds_list[0])
                image = st.session_state['cached_images'][analysis]
                st.empty().image(image)
            
            # Special case where the analysis output is audio player
            elif st.session_state['analysis_menu'][analysis] == 'listen':
                analysis_functions[analysis](sounds_list[0])

            # If the help button is pressed, the space to display analyses
            # is used to display help messages and images
            elif st.session_state['analysis_menu'][analysis] == 'help':
                helpstr = analysis_helps[analysis]
                st.empty().markdown(helpstr)
                st.empty().image(analysis_figures[analysis], use_column_width=True)

        # Columns for the menu below the analyses
        colrestart, colrep, coldown = st.columns(3)
        # Manually restart the analyses (clears all figures)
        colrestart.button(label='Réinitialiser les analyses',
                          on_click=reset_analyses,
                          args=(analysis_names, ))

        # Generate a report document using the performed analyses
        if colrep.button(label='Générer un rapport',
                         on_click=generate_report,
                         args=(st.session_state['report_analyses'],
                               st.session_state['cached_images'])):
            # After pressing the "generate" button a download button appears
            coldown.download_button(label='Télécharger le rapport',
                                    data=st.session_state['report_file'].read(),
                                    file_name='report.docx')

    # Case where no sounds were uploaded
    else:
        # Case when sounds were previously uploaded and then removed
        try:
            _ = analysis_names
        # This is reached if no sounds were uploaded
        except:
            # Default to displaying single sound analyses
            analysis_names = defined_analyses.single_sound_analysis_names
        reset_analyses(analysis_names)
        st.write('Veuillez sélectioner un son')

# A tab displaying information about the application
with help_tab:
    with open(os.path.join('documentation', 'documentation.md')) as md_file:
        md_string = md_file.read()
    st.markdown(md_string)

# A tab displaying information about the organisation
with about:
    with open(os.path.join('documentation', 'about.md')) as abt_file:
        md_string = abt_file.read()
    st.markdown(md_string)

